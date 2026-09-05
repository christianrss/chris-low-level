"""Build Treino_LowLevel_Unificado DOCX from day markdown modules.

Optional export only — not a quality gate. The authoritative study format is
modular Markdown under each day module folder (see docs/PEDAGOGY_STANDARD.md).
"""
from __future__ import annotations

import argparse
import re
import subprocess
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parents[1]

MODULE_ORDER_HINTS = [
    "linux", "systems", "ai", "algorithms", "quantum", "os", "debugger",
    "redteam", "dotnet", "nodejs", "javascript", "graphics", "terminal",
    "network", "p2p", "blockchain", "hardware", "tooling", "assembly",
    "architecture", "boot",
]

SECTION_FILES = [
    "README.md",
    "TEORIA_PASSO_A_PASSO.md",
    "PESQUISA_GUIADA.md",
    "EXERCICIOS.md",
    "RESOLUCAO_GUIADA_PASSO_A_PASSO.md",
    "RESOLUCAO_APENDICE.md",
    "TESTES_GUIADOS.md",
    "BENCHMARK_GUIADO.md",
]


def find_modules(day_dir: Path) -> list[Path]:
    modules = sorted(p.parent for p in day_dir.glob("*/*/RESOLUCAO_GUIADA_PASSO_A_PASSO.md"))
    def sort_key(m: Path) -> tuple[int, str]:
        track = m.parent.name
        try:
            idx = MODULE_ORDER_HINTS.index(track)
        except ValueError:
            idx = 99
        return (idx, m.name)
    return sorted(modules, key=sort_key)


def module_title(module: Path) -> str:
    readme = module / "README.md"
    if readme.exists():
        first = readme.read_text(encoding="utf-8").splitlines()
        for line in first:
            if line.startswith("#"):
                return line.lstrip("#").strip()
    return f"{module.parent.name}/{module.name}"


def stitch_day_markdown(day_dir: Path) -> str:
    date = day_dir.name
    parts: list[str] = [
        f"# Treino Low-Level Unificado — {date}\n",
        f"*Gerado automaticamente a partir dos Markdown em `{day_dir.relative_to(ROOT)}`.*\n",
    ]
    day_readme = day_dir / "README.md"
    if day_readme.exists():
        parts.append("## Visão geral do dia\n")
        parts.append(day_readme.read_text(encoding="utf-8"))
        parts.append("\n")

    start = day_dir / "START_HERE.md"
    if start.exists():
        parts.append("## Como começar\n")
        parts.append(start.read_text(encoding="utf-8"))
        parts.append("\n")

    for i, module in enumerate(find_modules(day_dir), 1):
        parts.append(f"\n---\n\n# Módulo {i}: {module_title(module)}\n")
        parts.append(f"`{module.relative_to(ROOT).as_posix()}`\n\n")
        for fname in SECTION_FILES:
            fpath = module / fname
            if fpath.exists():
                parts.append(f"## {fname.replace('.md', '').replace('_', ' ')}\n\n")
                parts.append(fpath.read_text(encoding="utf-8"))
                parts.append("\n")

    todo_map = day_dir / "TODO_MAP.md"
    if todo_map.exists():
        parts.append("\n---\n\n# Mapa global de TODOs\n\n")
        parts.append(todo_map.read_text(encoding="utf-8"))

    validation = day_dir / "VALIDATION.md"
    if validation.exists():
        parts.append("\n---\n\n# Validação\n\n")
        parts.append(validation.read_text(encoding="utf-8"))

    return "\n".join(parts)


def build_with_pandoc(md_path: Path, docx_path: Path) -> bool:
    try:
        subprocess.run(
            ["pandoc", str(md_path), "-o", str(docx_path), "--from", "markdown", "--to", "docx"],
            check=True,
            capture_output=True,
            text=True,
        )
        return True
    except (subprocess.CalledProcessError, FileNotFoundError) as exc:
        print(f"pandoc failed: {exc}", file=sys.stderr)
        return False


def build_with_python(md_path: Path, docx_path: Path) -> bool:
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed; install with: pip install python-docx", file=sys.stderr)
        return False

    doc = Document()
    content = md_path.read_text(encoding="utf-8")
    in_code = False
    code_lines: list[str] = []

    for line in content.splitlines():
        if line.strip().startswith("```"):
            if in_code:
                para = doc.add_paragraph("\n".join(code_lines))
                para.style = "No Spacing"
                code_lines = []
                in_code = False
            else:
                in_code = True
            continue
        if in_code:
            code_lines.append(line)
            continue
        if line.startswith("# "):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
        elif line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith("|") and "|" in line[1:]:
            doc.add_paragraph(line)
        elif line.strip():
            doc.add_paragraph(line)
        else:
            doc.add_paragraph("")

    doc.save(str(docx_path))
    return True


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--day", required=True, help="e.g. 2026-09-03")
    args = parser.parse_args()

    day_dir = ROOT / "days" / args.day
    if not day_dir.exists():
        print(f"Day folder not found: {day_dir}", file=sys.stderr)
        return 1

    stitched = stitch_day_markdown(day_dir)
    build_dir = ROOT / "build" / "docx"
    build_dir.mkdir(parents=True, exist_ok=True)
    md_path = build_dir / f"Treino_LowLevel_Unificado_{args.day}.md"
    docx_path = day_dir / f"Treino_LowLevel_Unificado_{args.day}.docx"
    md_path.write_text(stitched, encoding="utf-8")

    if build_with_pandoc(md_path, docx_path):
        print(f"DOCX written via pandoc: {docx_path}")
        return 0
    if build_with_python(md_path, docx_path):
        print(f"DOCX written via python-docx: {docx_path}")
        return 0
    print("Failed to build DOCX", file=sys.stderr)
    return 1


if __name__ == "__main__":
    raise SystemExit(main())
